import io
from docx import Document
from .base import BaseProcessor, ExtractionResult


class DocxProcessor(BaseProcessor):
    async def extract(self, file_content: bytes, filename: str) -> ExtractionResult:
        doc = Document(io.BytesIO(file_content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text for cell in row.cells])
            if rows:
                tables.append({"headers": rows[0], "rows": rows[1:]})

        return ExtractionResult(
            text="\n".join(paragraphs),
            tables=tables,
            metadata={"filename": filename, "paragraph_count": len(paragraphs)},
        )
